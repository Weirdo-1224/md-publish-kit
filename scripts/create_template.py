from __future__ import annotations

import argparse
import json
import shutil
import subprocess
import tempfile
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.shared import Cm, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from lxml import etree


def hex_rgb(value: str) -> RGBColor:
    return RGBColor.from_string(value.lstrip('#').upper())


def set_font(style, name: str, size: float | None = None, bold: bool | None = None, color: str | None = None):
    style.font.name = name
    rpr = style._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.insert(0, rfonts)
    for attr in ('ascii', 'hAnsi', 'eastAsia', 'cs'):
        rfonts.set(qn(f'w:{attr}'), name)
    if size is not None:
        style.font.size = Pt(size)
    if bold is not None:
        style.font.bold = bold
    if color:
        style.font.color.rgb = hex_rgb(color)


def set_run_font(run, name: str, size: float, color: str, bold: bool = False):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = hex_rgb(color)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.insert(0, rfonts)
    for attr in ('ascii', 'hAnsi', 'eastAsia', 'cs'):
        rfonts.set(qn(f'w:{attr}'), name)


def set_bottom_border(paragraph, color: str, size: str = '4'):
    ppr = paragraph._p.get_or_add_pPr()
    pbdr = ppr.find(qn('w:pBdr'))
    if pbdr is None:
        pbdr = OxmlElement('w:pBdr')
        ppr.append(pbdr)
    bottom = pbdr.find(qn('w:bottom'))
    if bottom is None:
        bottom = OxmlElement('w:bottom')
        pbdr.append(bottom)
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), size)
    bottom.set(qn('w:space'), '2')
    bottom.set(qn('w:color'), color)


def add_page_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement('w:fldChar')
    begin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = ' PAGE '
    end = OxmlElement('w:fldChar')
    end.set(qn('w:fldCharType'), 'end')
    run._r.extend([begin, instr, end])
    return run


def clear_container(container):
    for child in list(container._element):
        container._element.remove(child)


def patch_table_style(path: Path, cfg: dict):
    theme = cfg['theme']
    fonts = cfg['fonts']
    tmp = path.with_suffix('.patching.docx')
    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    with zipfile.ZipFile(path, 'r') as zin, zipfile.ZipFile(tmp, 'w', zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename == 'word/styles.xml':
                root = etree.fromstring(data)
                styles = root.xpath("//w:style[@w:styleId='Table']", namespaces=ns)
                if styles:
                    style = styles[0]
                    for node in list(style.xpath('./w:tblStylePr', namespaces=ns)):
                        style.remove(node)
                    tblpr = style.find(qn('w:tblPr'))
                    if tblpr is None:
                        tblpr = etree.SubElement(style, qn('w:tblPr'))
                    borders = tblpr.find(qn('w:tblBorders'))
                    if borders is None:
                        borders = etree.SubElement(tblpr, qn('w:tblBorders'))
                    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
                        el = borders.find(qn(f'w:{edge}'))
                        if el is None:
                            el = etree.SubElement(borders, qn(f'w:{edge}'))
                        el.set(qn('w:val'), 'single')
                        el.set(qn('w:sz'), '4')
                        el.set(qn('w:color'), theme['border'])
                    rpr = style.find(qn('w:rPr'))
                    if rpr is None:
                        rpr = etree.SubElement(style, qn('w:rPr'))
                    rfonts = rpr.find(qn('w:rFonts'))
                    if rfonts is None:
                        rfonts = etree.SubElement(rpr, qn('w:rFonts'))
                    for attr in ('ascii', 'hAnsi', 'eastAsia', 'cs'):
                        rfonts.set(qn(f'w:{attr}'), fonts['body'])
                    sz = rpr.find(qn('w:sz'))
                    if sz is None:
                        sz = etree.SubElement(rpr, qn('w:sz'))
                    sz.set(qn('w:val'), str(int(fonts['table_size_pt'] * 2)))
                data = etree.tostring(root, xml_declaration=True, encoding='UTF-8', standalone='yes')
            zout.writestr(item, data)
    tmp.replace(path)


def build_template(config_path: Path, output_path: Path):
    cfg = json.loads(config_path.read_text(encoding='utf-8'))
    doc_cfg, theme, fonts = cfg['document'], cfg['theme'], cfg['fonts']

    with tempfile.TemporaryDirectory() as td:
        base = Path(td) / 'reference.docx'
        with base.open('wb') as f:
            subprocess.run(['pandoc', '--print-default-data-file', 'reference.docx'], stdout=f, check=True)
        doc = Document(base)

    section = doc.sections[0]
    page = doc_cfg['page']
    section.page_width = Cm(page['width_cm'])
    section.page_height = Cm(page['height_cm'])
    section.top_margin = Cm(page['top_margin_cm'])
    section.bottom_margin = Cm(page['bottom_margin_cm'])
    section.left_margin = Cm(page['left_margin_cm'])
    section.right_margin = Cm(page['right_margin_cm'])
    section.header_distance = Cm(0.85)
    section.footer_distance = Cm(0.75)

    # Clean, single-line header.
    clear_container(section.header)
    hp = section.header.add_paragraph()
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = hp.add_run(doc_cfg['project_name'])
    set_run_font(r, fonts['body'], 9, theme['primary'], True)
    r = hp.add_run('  ·  ' + doc_cfg['header_text'])
    set_run_font(r, fonts['body'], 9, theme['muted'])
    set_bottom_border(hp, theme['border'])

    # Clean, single-line footer with a right-aligned page field.
    clear_container(section.footer)
    fp = section.footer.add_paragraph()
    usable_cm = page['width_cm'] - page['left_margin_cm'] - page['right_margin_cm']
    fp.paragraph_format.tab_stops.add_tab_stop(Cm(usable_cm), WD_TAB_ALIGNMENT.RIGHT)
    r = fp.add_run(doc_cfg['footer_text'])
    set_run_font(r, fonts['body'], 8, theme['muted'])
    fp.add_run('\t')
    r = fp.add_run('第 ')
    set_run_font(r, fonts['body'], 8, theme['muted'])
    page_run = add_page_field(fp)
    set_run_font(page_run, fonts['body'], 8, theme['muted'])
    r = fp.add_run(' 页')
    set_run_font(r, fonts['body'], 8, theme['muted'])

    styles = doc.styles
    normal = styles['Normal']
    set_font(normal, fonts['body'], fonts['body_size_pt'], False, theme['text'])
    normal.paragraph_format.line_spacing = 1.45
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.widow_control = True

    for name in ('Body Text', 'First Paragraph', 'Compact', 'Block Text'):
        if name in styles:
            set_font(styles[name], fonts['body'], fonts['body_size_pt'], False, theme['text'])
            styles[name].paragraph_format.line_spacing = 1.4
            styles[name].paragraph_format.space_after = Pt(4)

    title = styles['Title']
    set_font(title, fonts['body'], 26, True, theme['primary'])
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(88)
    title.paragraph_format.space_after = Pt(13)

    if 'Subtitle' in styles:
        subtitle = styles['Subtitle']
        set_font(subtitle, fonts['body'], 13, False, theme['muted'])
        subtitle.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.paragraph_format.space_after = Pt(24)
    if 'Author' in styles:
        author = styles['Author']
        set_font(author, fonts['body'], 11, False, theme['text'])
        author.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        author.paragraph_format.space_after = Pt(7)
    if 'Date' in styles:
        date = styles['Date']
        set_font(date, fonts['body'], 10, False, theme['muted'])
        date.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    heading_specs = {
        'Heading1': (17, theme['primary'], 18, 8),
        'Heading2': (13.5, theme['accent'], 14, 6),
        'Heading3': (11.5, theme['text'], 11, 5),
    }
    for name, (size, color, before, after) in heading_specs.items():
        st = styles[name]
        set_font(st, fonts['body'], size, True, color)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True

    if 'Source Code' not in styles:
        code = styles.add_style('Source Code', WD_STYLE_TYPE.PARAGRAPH)
    else:
        code = styles['Source Code']
    set_font(code, fonts['code'], fonts['code_size_pt'], False, theme['text'])
    code.paragraph_format.left_indent = Cm(0.35)
    code.paragraph_format.right_indent = Cm(0.35)
    code.paragraph_format.space_before = Pt(5)
    code.paragraph_format.space_after = Pt(6)
    code.paragraph_format.line_spacing = 1.15

    for name in ('Caption', 'Table Caption', 'Image Caption'):
        if name in styles:
            st = styles[name]
            set_font(st, fonts['body'], 9, False, theme['muted'])
            st.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            st.paragraph_format.space_before = Pt(4)
            st.paragraph_format.space_after = Pt(6)

    # Reference-doc body is ignored by Pandoc; remove it to keep the file clean.
    body = doc._element.body
    for child in list(body):
        if child.tag != qn('w:sectPr'):
            body.remove(child)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    patch_table_style(output_path, cfg)
    print(f'Created template: {output_path}')


def main():
    parser = argparse.ArgumentParser(description='Create the Compass Pandoc reference DOCX template.')
    parser.add_argument('--config', type=Path, required=True)
    parser.add_argument('--output', type=Path, required=True)
    args = parser.parse_args()
    if shutil.which('pandoc') is None:
        raise SystemExit('Pandoc not found. Install Pandoc before generating the template.')
    build_template(args.config.resolve(), args.output.resolve())


if __name__ == '__main__':
    main()

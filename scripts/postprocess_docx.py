from __future__ import annotations

import argparse
import json
import math
import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def rgb(value: str) -> RGBColor:
    return RGBColor.from_string(value.lstrip('#').upper())


def set_east_asia_font(run, name: str):
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.insert(0, rfonts)
    for attr in ('ascii', 'hAnsi', 'eastAsia', 'cs'):
        rfonts.set(qn(f'w:{attr}'), name)


def set_shading(cell_or_paragraph, fill: str):
    if hasattr(cell_or_paragraph, '_tc'):
        parent = cell_or_paragraph._tc.get_or_add_tcPr()
    else:
        parent = cell_or_paragraph._p.get_or_add_pPr()
    shd = parent.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        parent.append(shd)
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), fill)


def set_cell_margins(cell, value: int):
    tcpr = cell._tc.get_or_add_tcPr()
    tcmar = tcpr.find(qn('w:tcMar'))
    if tcmar is None:
        tcmar = OxmlElement('w:tcMar')
        tcpr.append(tcmar)
    for key in ('top', 'start', 'bottom', 'end'):
        node = tcmar.find(qn(f'w:{key}'))
        if node is None:
            node = OxmlElement(f'w:{key}')
            tcmar.append(node)
        node.set(qn('w:w'), str(value))
        node.set(qn('w:type'), 'dxa')


def set_cell_borders(cell, color: str):
    tcpr = cell._tc.get_or_add_tcPr()
    borders = tcpr.find(qn('w:tcBorders'))
    if borders is None:
        borders = OxmlElement('w:tcBorders')
        tcpr.append(borders)
    for edge in ('top', 'left', 'bottom', 'right'):
        node = borders.find(qn(f'w:{edge}'))
        if node is None:
            node = OxmlElement(f'w:{edge}')
            borders.append(node)
        node.set(qn('w:val'), 'single')
        node.set(qn('w:sz'), '4')
        node.set(qn('w:color'), color)


def set_repeat_header(row):
    trpr = row._tr.get_or_add_trPr()
    node = trpr.find(qn('w:tblHeader'))
    if node is None:
        node = OxmlElement('w:tblHeader')
        trpr.append(node)
    node.set(qn('w:val'), 'true')


def prevent_row_split(row):
    trpr = row._tr.get_or_add_trPr()
    node = trpr.find(qn('w:cantSplit'))
    if node is None:
        node = OxmlElement('w:cantSplit')
        trpr.append(node)


def set_heading_border(paragraph, color: str):
    ppr = paragraph._p.get_or_add_pPr()
    pbdr = ppr.find(qn('w:pBdr'))
    if pbdr is None:
        pbdr = OxmlElement('w:pBdr')
        ppr.append(pbdr)
    bottom = pbdr.find(qn('w:bottom'))
    if bottom is None:
        bottom = OxmlElement('w:bottom')
        pbdr.append(bottom)
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), color)


def normalize_header(text: str) -> str:
    return re.sub(r'\s+', '', text).strip().lower()


def find_rule(headers: list[str], rules: list[dict]) -> dict | None:
    normalized = [normalize_header(x) for x in headers]
    for rule in rules:
        if normalized == [normalize_header(x) for x in rule['headers']]:
            return rule
    return None


def char_weight(text: str) -> float:
    # CJK characters are visually wider; URLs/code need extra width.
    cjk = sum(1 for ch in text if '\u4e00' <= ch <= '\u9fff')
    ascii_count = len(text) - cjk
    score = cjk * 1.8 + ascii_count
    if 'http' in text or '/' in text or '_' in text:
        score *= 1.15
    return max(score, 4.0)


def infer_ratios(table) -> list[float]:
    cols = len(table.columns)
    scores = []
    for j in range(cols):
        values = [cell.text.strip() for row in table.rows for cell in [row.cells[j]]]
        lengths = sorted((char_weight(v) for v in values), reverse=True)
        representative = sum(lengths[: min(4, len(lengths))]) / max(1, min(4, len(lengths)))
        scores.append(max(representative, 6.0))

    # First columns containing sequence/method labels should stay compact.
    headers = [normalize_header(c.text) for c in table.rows[0].cells]
    for idx, h in enumerate(headers):
        if h in {'#', '序号', '方法', '版本'}:
            scores[idx] *= 0.55
    total = sum(scores)
    ratios = [s / total for s in scores]

    # Enforce reasonable minimums and maximums, then renormalize.
    minimum = 0.08 if cols >= 4 else 0.12
    ratios = [max(minimum, min(0.62, x)) for x in ratios]
    total = sum(ratios)
    return [x / total for x in ratios]


def alignment_from_text(header: str, body_text: str) -> str:
    h = normalize_header(header)
    if h in {'#', '序号', '姓名', '版本', '方法', '中文名', '输入', '输出', '角色', '模块', '层级', '模型', '提供商', '协议', '维度', '表名'}:
        return 'center'
    if len(body_text.strip()) <= 14 and '\n' not in body_text:
        return 'center'
    return 'left'


def align_paragraph(paragraph, mode: str):
    paragraph.alignment = {
        'center': WD_ALIGN_PARAGRAPH.CENTER,
        'right': WD_ALIGN_PARAGRAPH.RIGHT,
        'left': WD_ALIGN_PARAGRAPH.LEFT,
    }.get(mode, WD_ALIGN_PARAGRAPH.LEFT)


def set_table_widths(table, usable_width_emu: int, ratios: list[float]):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for idx, ratio in enumerate(ratios):
        width = int(usable_width_emu * ratio)
        table.columns[idx].width = width
        for cell in table.columns[idx].cells:
            cell.width = width
            tcpr = cell._tc.get_or_add_tcPr()
            tcw = tcpr.find(qn('w:tcW'))
            if tcw is None:
                tcw = OxmlElement('w:tcW')
                tcpr.append(tcw)
            tcw.set(qn('w:w'), str(int(width / 635)))  # EMU -> twips approximately
            tcw.set(qn('w:type'), 'dxa')
    grid = table._tbl.tblGrid
    if grid is not None:
        grid_cols = list(grid.iterchildren())
        for idx, ratio in enumerate(ratios):
            if idx < len(grid_cols):
                grid_cols[idx].set(qn('w:w'), str(int((usable_width_emu * ratio) / 635)))


def style_tables(doc, cfg: dict):
    theme, fonts = cfg['theme'], cfg['fonts']
    table_cfg = cfg['tables']
    section = doc.sections[0]
    usable_width = section.page_width - section.left_margin - section.right_margin

    for table in doc.tables:
        if not table.rows or not table.columns:
            continue
        headers = [cell.text.strip() for cell in table.rows[0].cells]
        rule = find_rule(headers, table_cfg['rules'])
        ratios = rule['ratios'] if rule else infer_ratios(table)
        aligns = rule.get('align', []) if rule else []
        if len(ratios) != len(table.columns):
            ratios = infer_ratios(table)
        set_table_widths(table, usable_width, ratios)
        set_repeat_header(table.rows[0])

        for row_idx, row in enumerate(table.rows):
            prevent_row_split(row)
            for col_idx, cell in enumerate(row.cells):
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                set_cell_margins(cell, int(table_cfg['cell_margin_dxa']))
                set_cell_borders(cell, theme['border'])
                if row_idx == 0:
                    set_shading(cell, theme['primary'])
                elif row_idx % 2 == 0:
                    set_shading(cell, theme['band'])
                else:
                    set_shading(cell, theme['white'])

                body_text = cell.text.strip()
                mode = 'center' if row_idx == 0 else (
                    aligns[col_idx] if col_idx < len(aligns) else alignment_from_text(headers[col_idx], body_text)
                )
                for paragraph in cell.paragraphs:
                    align_paragraph(paragraph, mode)
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(0)
                    paragraph.paragraph_format.line_spacing = 1.2
                    for run in paragraph.runs:
                        set_east_asia_font(run, fonts['body'])
                        run.font.size = Pt(fonts['table_size_pt'])
                        run.font.color.rgb = rgb(theme['white'] if row_idx == 0 else theme['text'])
                        if row_idx == 0:
                            run.bold = True


def style_paragraphs(doc, cfg: dict):
    theme, fonts = cfg['theme'], cfg['fonts']
    has_cover = any(p.style and p.style.name == 'Title' and p.text.strip() for p in doc.paragraphs)
    first_h1 = True
    for paragraph in doc.paragraphs:
        style_name = paragraph.style.name if paragraph.style else ''
        if style_name == 'Heading 1' or style_name == 'Heading1':
            set_heading_border(paragraph, theme['accent'])
            if first_h1 and has_cover and cfg['document'].get('cover_page', True):
                paragraph.paragraph_format.page_break_before = True
            first_h1 = False
        if style_name == 'Source Code':
            set_shading(paragraph, theme['code_bg'])
            paragraph.paragraph_format.left_indent = Cm(0.35)
            paragraph.paragraph_format.right_indent = Cm(0.35)
            for run in paragraph.runs:
                set_east_asia_font(run, fonts['code'])
                run.font.size = Pt(fonts['code_size_pt'])
        elif style_name in {'Block Text', 'Intense Quote', 'Quote'}:
            set_shading(paragraph, theme['light'])
            paragraph.paragraph_format.left_indent = Cm(0.35)
            paragraph.paragraph_format.right_indent = Cm(0.2)
        for run in paragraph.runs:
            if style_name != 'Source Code':
                set_east_asia_font(run, fonts['body'])


def style_images(doc, cfg: dict):
    section = doc.sections[0]
    usable = section.page_width - section.left_margin - section.right_margin
    max_width = int(usable * cfg['images']['max_width_ratio'])
    for shape in doc.inline_shapes:
        if shape.width > max_width:
            ratio = max_width / shape.width
            shape.width = max_width
            shape.height = int(shape.height * ratio)
        if cfg['images'].get('center', True):
            inline = shape._inline
            paragraph_el = inline.getparent().getparent()
            for paragraph in doc.paragraphs:
                if paragraph._p is paragraph_el:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    break


def postprocess(input_path: Path, output_path: Path, config_path: Path):
    cfg = json.loads(config_path.read_text(encoding='utf-8'))
    doc = Document(input_path)
    style_paragraphs(doc, cfg)
    style_tables(doc, cfg)
    style_images(doc, cfg)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    print(f'Postprocessed: {output_path}')


def main():
    parser = argparse.ArgumentParser(description='Normalize Pandoc DOCX for Feishu import.')
    parser.add_argument('input', type=Path)
    parser.add_argument('output', type=Path)
    parser.add_argument('--config', type=Path, required=True)
    args = parser.parse_args()
    postprocess(args.input.resolve(), args.output.resolve(), args.config.resolve())


if __name__ == '__main__':
    main()
